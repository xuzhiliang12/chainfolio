from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Chainfolio_使用说明书.docx"

# compact_reference_guide preset, with one named brand override:
# lime #A7FF4F and cyan #4AD9E8 are used only for product-identifying accents.
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6775"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
LIME = "A7FF4F"
CYAN = "4AD9E8"
GREEN_DARK = "315B22"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, color=INK, bold=None, italic=None, latin=FONT_LATIN, cjk=FONT_CJK):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(p, before=0, after=6, line=1.25, keep_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.widow_control = True


def add_text(doc, text, *, size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6, line=1.25):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, line)
    if align is not None:
        p.alignment = align
    set_run_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_paragraph(doc, parts, *, before=0, after=6, line=1.25):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, line)
    for text, bold, color in parts:
        set_run_font(p.add_run(text), size=11, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_kicker(doc, text):
    return add_text(doc, text.upper(), size=9, color=GREEN_DARK, bold=True, after=4)


def add_callout(doc, title, body, tone="info"):
    colors = {
        "info": (LIGHT_BLUE, DARK_BLUE),
        "safe": ("EEFAE7", GREEN_DARK),
        "warn": ("FFF6D9", GOLD),
        "risk": ("FDECEC", RED),
    }
    fill, color = colors[tone]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    style_paragraph(p, 0, 0, 1.15)
    set_run_font(p.add_run(title + "  "), size=10.5, color=color, bold=True)
    set_run_font(p.add_run(body), size=10.5, color=INK)
    add_text(doc, "", after=2)
    return table


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, 0, 0, 1.05)
        set_run_font(p.add_run(header), size=font_size, color=INK, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            style_paragraph(p, 0, 0, 1.08)
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    add_text(doc, "", after=2)
    return table


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(document):
    numbering = document.part.numbering_part.element

    def make(abstract_id, num_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, jc, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)

    make(50, 50, "bullet", "•", FONT_LATIN)
    make(51, 51, "decimal", "%1.")
    return 50, 51


def add_list_item(doc, text, num_id, bold_prefix=None):
    actual_num_id = num_id
    if num_id == 51:
        ordered_ids = getattr(doc, "_chainfolio_ordered_ids", set())
        previous_num_id = None
        if doc.paragraphs:
            prev_num = doc.paragraphs[-1]._p.find("./w:pPr/w:numPr/w:numId", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if prev_num is not None:
                previous_num_id = int(prev_num.get(qn("w:val")))
        if previous_num_id in ordered_ids:
            actual_num_id = previous_num_id
        else:
            numbering = doc.part.numbering_part.element
            existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
            actual_num_id = max(existing, default=51) + 1
            abstracts = numbering.findall(qn("w:abstractNum"))
            abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in abstracts]
            actual_abstract_id = max(abstract_ids, default=51) + 1
            source_abstract = next(node for node in abstracts if node.get(qn("w:abstractNumId")) == "51")
            cloned_abstract = copy.deepcopy(source_abstract)
            cloned_abstract.set(qn("w:abstractNumId"), str(actual_abstract_id))
            first_num = numbering.find(qn("w:num"))
            numbering.insert(list(numbering).index(first_num), cloned_abstract)
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(actual_num_id))
            abstract_num_id = OxmlElement("w:abstractNumId")
            abstract_num_id.set(qn("w:val"), str(actual_abstract_id))
            num.append(abstract_num_id)
            lvl_override = OxmlElement("w:lvlOverride")
            lvl_override.set(qn("w:ilvl"), "0")
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            lvl_override.append(start_override)
            num.append(lvl_override)
            numbering.append(num)
            ordered_ids.add(actual_num_id)
            doc._chainfolio_ordered_ids = ordered_ids
    p = doc.add_paragraph()
    style_paragraph(p, 0, 4, 1.25)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(actual_num_id))
    num_pr.extend([ilvl, num])
    p_pr.insert(0, num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=11, color=INK, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=11, color=INK)
    else:
        set_run_font(p.add_run(text), size=11, color=INK)
    return p


def add_page(doc, title, kicker=None, subtitle=None):
    doc.add_page_break()
    if kicker:
        add_kicker(doc, kicker)
    add_heading(doc, title, 1)
    if subtitle:
        add_text(doc, subtitle, size=10.5, color=MUTED, italic=True, after=10)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = doc.styles[name]
        st.font.name = FONT_LATIN
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(hp, 0, 0, 1.0)
    set_run_font(hp.add_run("CHAINFOLIO  |  使用说明书"), size=9, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(fp, 0, 0, 1.0)
    set_run_font(fp.add_run("Chainfolio · 第 "), size=9, color=MUTED)
    page_field(fp)
    set_run_font(fp.add_run(" 页"), size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)
    bullet_id, number_id = create_numbering(doc)
    props = doc.core_properties
    props.title = "Chainfolio 使用说明书"
    props.subject = "多链只读资产总账的功能介绍与操作指南"
    props.author = "Chainfolio"
    props.keywords = "Chainfolio, 多链资产, 钱包管理, 使用说明"

    # Cover: editorial_cover pattern.
    add_text(doc, "CHAINFOLIO", size=11, color=GREEN_DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=70, after=20)
    add_text(doc, "使用说明书", size=30, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_text(doc, "多人管理 · 多链地址 · 只读资产总账", size=15, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_text(doc, "负责人 → 手机 → OKX 钱包 → 链上地址 → 资产", size=11.5, color=GREEN_DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=60)
    add_text(doc, "适用对象", size=9, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "管理员、资产负责人及日常运营人员", size=11, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_text(doc, "版本 1.0  ·  2026 年 8 月", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, 0, 0, 1.0)
    add_hyperlink(p, "https://github.com/xuzhiliang12/chainfolio", "github.com/xuzhiliang12/chainfolio")

    add_page(doc, "1. 快速上手", "QUICK START", "按顺序完成以下步骤，即可建立第一套可用的资产总账。")
    for text in (
        "首次打开时创建自己的本地主账户；已经初始化的设备直接登录。",
        "在“负责人”中添加实际管理人，例如张三、李四。",
        "在“手机”中添加设备并指定负责人；系统默认建立 3 个钱包，可增加到 10 个。",
        "在“链上地址”中选择手机和钱包，然后导入 EVM 或 Solana 公共地址。",
        "点击“手动更新”读取余额，随后在“资产总览”和“资产明细”查看汇总结果。",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "最重要的规则", "本系统是 WATCH ONLY（只读监控）。只输入公开地址，绝不要输入助记词、私钥或 OKX 密码。", "safe")
    add_heading(doc, "推荐的首次配置顺序", 2)
    add_table(doc, ["顺序", "配置对象", "完成标志"], [
        ("1", "负责人", "每组手机都有明确归属"),
        ("2", "手机", "名称易辨认，负责人正确"),
        ("3", "钱包", "名称、生成时间与状态已登记"),
        ("4", "链上地址", "地址属于正确的手机和钱包"),
        ("5", "币种 / 链", "按需补充自定义币种或 EVM 链"),
    ], [1100, 2500, 5760])

    add_page(doc, "2. 系统框架与菜单", "SYSTEM MODEL", "先理解数据归属关系，再开始批量导入。")
    add_heading(doc, "数据层级", 2)
    add_callout(doc, "一条清晰的归属链", "系统用户 → 负责人 → 手机 → OKX 钱包 → 链上地址 → 资产。下级数据会逐层汇总到上级，最后形成当前用户的总账户数据。", "info")
    add_table(doc, ["层级", "用途"], [
        ("系统用户", "登录账号；不同用户的数据互相隔离"),
        ("负责人", "统计某个人管理的全部手机、钱包与资产"),
        ("手机", "记录设备；必须归属一位负责人"),
        ("OKX 钱包", "记录生成时间、女巫检查、活动进度和备注"),
        ("链上地址", "公开地址；必须归属某台手机的某个钱包"),
        ("资产", "从链上读取的币种余额与美元价值"),
    ], [2700, 6660])
    add_heading(doc, "左侧菜单", 2)
    add_table(doc, ["分组", "页面", "主要用途"], [
        ("总览", "资产总览", "净值、时间对比、链/币种分布、风险与活动摘要"),
        ("总览", "钱包管理", "像 Excel 一样按钱包统计、筛选、批量更新与导出"),
        ("资产管理", "负责人 / 手机", "维护组织层级；从手机中展开和维护钱包资料"),
        ("资产管理", "链上地址 / 资产明细", "导入地址、查看币种余额与价值"),
        ("系统配置", "链配置", "查看默认网络并导入自定义 EVM 链"),
        ("系统配置", "用户管理", "主账户按需开启团队注册、生成邀请码并管理用户"),
    ], [1600, 2500, 5260])

    add_page(doc, "3. 主账户、登录与团队模式", "ACCESS", "每台自托管设备拥有独立的主账户和数据空间。")
    add_heading(doc, "首次创建本地主账户", 2)
    for text in ("打开 http://127.0.0.1:4173。", "设置本地主账户用户名和不少于 10 位的密码。", "创建完成后初始化入口自动关闭，并直接进入资产总览。"):
        add_list_item(doc, text, number_id)
    add_heading(doc, "登录", 2)
    for text in ("打开系统网址。", "输入用户名和密码。", "点击“登录”进入资产总览。"):
        add_list_item(doc, text, number_id)
    add_heading(doc, "可选的团队邀请码注册", 2)
    for text in ("主账户在“用户管理”中开启团队注册。", "生成一次性邀请码并发送给指定成员。", "成员切换到“团队注册”，设置独立用户名和密码。", "单机个人使用时保持关闭，不需要任何邀请码。"):
        add_list_item(doc, text, number_id)
    add_callout(doc, "邀请码规则", "团队注册默认关闭。邀请码只能使用一次；主账户生成后应安全地发送给指定人员，不要公开转发。", "warn")
    add_heading(doc, "修改密码与退出", 2)
    add_text(doc, "点击右上角头像或用户名打开“账户设置”，输入当前密码和新密码后保存。修改成功后，其他设备上的旧登录会立即失效。需要结束当前会话时选择“退出登录”。")
    add_heading(doc, "显示偏好", 2)
    add_table(doc, ["设置", "可选项", "说明"], [
        ("字体大小", "90% / 100% / 115% / 130%", "选择后保存在当前设备"),
        ("隐私模式", "显示 / 隐藏", "隐藏页面中的余额、数量和美元价值"),
    ], [1800, 2900, 4660])

    add_page(doc, "4. 负责人和手机管理", "OWNERS & DEVICES", "负责人和手机是所有钱包数据的归属基础。")
    add_heading(doc, "添加负责人", 2)
    for text in ("进入“负责人”。", "点击“添加负责人”。", "输入姓名或内部识别名称并保存。"):
        add_list_item(doc, text, number_id)
    add_text(doc, "负责人卡片会汇总其名下的手机、钱包、地址和资产。负责人名称可以修改；删除前需要先处理其名下的手机归属。")
    add_heading(doc, "添加手机", 2)
    for text in ("进入“手机”，点击“添加手机”。", "输入设备名称，例如“工作手机 A”或“P-01”。", "选择负责人并保存。"):
        add_list_item(doc, text, number_id)
    add_callout(doc, "钱包数量", "每台手机默认创建 3 个 OKX 钱包；后续可按实际情况增加，单台手机最多 10 个钱包，并不要求固定为 3 个。", "info")
    add_heading(doc, "查看与维护", 2)
    for text in (
        "可按负责人筛选或搜索手机名称。",
        "负责人下拉框可直接调整手机归属。",
        "“查看钱包”会展开该手机下的全部钱包资料。",
        "长列表默认显示 20 条，可继续展开或折叠。",
    ):
        add_list_item(doc, text, bullet_id)

    add_page(doc, "5. 钱包资料、女巫检查与活动", "WALLET PROFILE", "这些字段均为内部人工记录，不会触发任何链上操作。")
    add_heading(doc, "编辑钱包资料", 2)
    add_text(doc, "在“手机”中展开钱包，或进入“钱包”页面，点击“编辑”。可维护钱包名称、生成时间、女巫检查状态、活动进度以及判断依据/备注。抽屉底部的“上一个 / 下一个”适合连续登记。")
    add_table(doc, ["女巫状态", "建议使用场景"], [
        ("未登记", "尚未录入检查结果"),
        ("待检查", "已进入检查流程，尚无结论"),
        ("未发现风险", "当前检查未发现明显风险"),
        ("疑似女巫", "存在相似行为等线索，需要复核"),
        ("已确认女巫", "内部已确认并记录依据"),
    ], [2500, 6860])
    add_heading(doc, "建立通用活动", 2)
    for text in (
        "打开“钱包台账”，点击“活动管理”。",
        "输入活动或任务名称，例如“OKX Cryptopedia 第 5 期”。",
        "保存后，该活动会自动成为所有钱包的一列，无需逐个钱包创建。",
        "在每个钱包中分别记录：未登记、未参加、已参加、进行中或已完成。",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "归档不会丢数据", "活动结束后可归档；历史进度会保留，并可随时恢复。当前账户最多创建 30 个通用活动。", "safe")

    add_page(doc, "6. 导入链上地址", "ADDRESS REGISTRY", "地址必须明确归属到具体手机和具体钱包。")
    add_heading(doc, "导入步骤", 2)
    for text in (
        "点击页面右上角“添加链上地址”，或进入“链上地址”后添加。",
        "选择手机。",
        "选择该手机中的 OKX 钱包。",
        "选择地址类型：EVM 多链或 Solana。",
        "粘贴公开地址，点击“添加并更新”。",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "EVM 地址只需导入一次", "同一个 EVM 地址会自动用于 Ethereum、BNB Chain、Base、Arbitrum、Optimism、Robinhood Chain、X Layer，以及以后导入的自定义 EVM 链。", "safe")
    add_heading(doc, "默认支持的区块链", 2)
    add_table(doc, ["地址类型", "网络"], [
        ("EVM 多链", "Ethereum、Arbitrum、Base、Optimism、BNB Chain、Robinhood Chain、X Layer"),
        ("Solana", "Solana 主网"),
    ], [2100, 7260])
    add_heading(doc, "修改或删除地址", 2)
    add_text(doc, "在“链上地址”列表中使用“编辑”修改归属或地址，使用“删除”移除监控记录。删除会影响对应资产汇总，操作前请核对目标。")

    add_page(doc, "7. 资产总览与净值对比", "PORTFOLIO", "总览用于快速判断总资产、变化和资产结构。")
    add_heading(doc, "总净值和时间对比", 2)
    add_text(doc, "资产总净值会汇总当前用户下的全部负责人、手机、钱包和地址。默认与昨日进行对比，也可选择 7 天、30 天、90 天或自定义日期，查看净值增加/减少金额和比例。趋势柱图用于辅助观察变化。")
    add_heading(doc, "资产分布", 2)
    add_table(doc, ["模式", "显示内容"], [
        ("按链", "每条链的资产价值、币种数量和占比"),
        ("按币种", "每种币的总数量、美元价值和占比"),
    ], [2100, 7260])
    add_text(doc, "点击“按链 / 按币种”切换口径。圆环中间显示链或币种数量，右侧图例显示明细。没有价格的数据不会形成有效美元分布。")
    add_heading(doc, "隐私模式", 2)
    add_text(doc, "点击右上角眼睛图标可隐藏余额。总净值、图表提示、分布美元价值、币种数量、负责人/手机/钱包净值、台账净值、资产持有量和价值都会被隐藏；名称、链、地址数量和百分比仍保留。")
    add_callout(doc, "导出提醒", "即使隐私模式已开启，钱包台账导出的 CSV 仍包含真实资产净值；系统会在导出前提醒。", "warn")

    add_page(doc, "8. 资产明细与自定义币种", "TOKENS", "用于补充系统没有自动识别的非主流币种。")
    add_heading(doc, "查看资产明细", 2)
    add_text(doc, "“资产明细”显示币种、所在链、负责人/手机/钱包来源、持有量、24H 变化和美元价值。可按链筛选、搜索币种或链，并隐藏小额资产。列表默认显示 20 条。")
    add_heading(doc, "添加自定义币种", 2)
    for text in (
        "点击“添加币种”，选择所在链。",
        "输入合约地址或 Solana Mint 地址，点击“自动识别”。",
        "核对系统读取的名称、符号和精度。非标准币种可手动修正。",
        "单价 USD 可留空自动报价；需要固定口径时也可手动填写。",
        "点击“添加并同步”。系统会扫描当前账户在该链上的全部地址。",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "一次添加，全账户生效", "自定义币种按当前系统用户的账户级配置保存。添加一次后，已有和以后新增的匹配地址都会自动纳入，不需要逐个钱包重复添加。", "safe")
    add_heading(doc, "自动报价规则", 2)
    add_text(doc, "系统优先选取流动性较高的 DEX 美元报价。若找不到有效交易对、流动性过低或价格源暂不可用，币种余额仍会显示，但价值会显示为“暂无报价”或 $0.00。此时可编辑币种并填写手动单价。")

    add_page(doc, "9. 自定义 EVM 链", "NETWORKS", "当项目链未预置时，可自行导入兼容 EVM 的网络。")
    add_heading(doc, "导入步骤", 2)
    for text in (
        "进入“链配置”，点击“导入自定义链”。",
        "填写网络名称。",
        "填写原生币符号。",
        "填写正确的 Chain ID。",
        "填写以 https:// 开头的可信 RPC URL 并保存。",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "地址自动复用", "导入新的 EVM 链后，系统会复用已经登记的 EVM 地址，无需再次逐个钱包添加地址。", "info")
    add_heading(doc, "配置建议", 2)
    for text in (
        "优先使用项目官方或可信服务商提供的 HTTPS RPC。",
        "确认 Chain ID 与网络一致，避免查询到错误网络。",
        "公开 RPC 可能有限流；批量查询较多时建议使用稳定的专用 RPC。",
        "自定义链的代币仍需通过“自定义币种”添加合约后才能完整统计。",
    ):
        add_list_item(doc, text, bullet_id)

    add_page(doc, "10. 钱包台账与统计", "WALLET OPERATIONS", "钱包台账承担 Excel 式的登记、统计和批量处理。")
    add_heading(doc, "统计区域", 2)
    add_table(doc, ["区域", "可查看内容"], [
        ("状态卡片", "钱包总数、已登记、待检查、正常、疑似、已确认"),
        ("负责人汇总", "按负责人统计无风险、疑似、确认和未检查数量"),
        ("活动完成情况", "已完成、已参加、进行中、未参加、未登记和完成率"),
        ("钱包明细矩阵", "一行一个钱包，活动相当于新增一列"),
    ], [2600, 6760])
    add_heading(doc, "筛选和批量操作", 2)
    for text in (
        "按负责人、手机、女巫状态筛选，或搜索钱包/手机。",
        "点击统计卡片可快速筛选对应状态。",
        "使用“显示列”决定是否展示生成时间、活动、地址数和净值。",
        "选择舒适或紧凑密度；长列表每页/每段默认 20 条。",
        "勾选当前页钱包，批量修改女巫检查或某项活动进度。",
        "点击“导出 CSV”生成可在 Excel 中继续分析的明细。",
    ):
        add_list_item(doc, text, bullet_id)

    add_page(doc, "11. 更新机制", "SYNC", "系统采用分散、随机的只读查询，避免固定频率集中请求。")
    add_heading(doc, "自动更新", 2)
    add_table(doc, ["规则", "说明"], [
        ("地址周期", "每个地址独立随机安排在约 10–24 小时后更新"),
        ("后台检查", "系统以随机间隔检查到期任务"),
        ("批次顺序", "到期地址会打乱顺序，并采用随机批次执行"),
        ("重新安排", "手动更新完成后，所选地址会重新生成下一次随机时间"),
    ], [2600, 6760])
    add_heading(doc, "手动更新", 2)
    add_text(doc, "点击“手动更新”，可选择全部账户、指定负责人、指定手机或指定钱包。只更新正在核对的范围，速度更快，也能减少不必要的 RPC 请求。更新过程只读取公开链上数据，不会签名或发送交易。")
    add_callout(doc, "重要说明", "随机刷新用于分散 RPC/API 负载和避免固定轮询节奏，但不能保证改变任何项目方的风控、反女巫或关联分析结果。钱包是否被判定为女巫，取决于项目规则和链上行为等多种因素。", "warn")
    add_heading(doc, "什么时候手动更新", 2)
    for text in ("刚导入地址或自定义币种后。", "刚完成转账，需要立即核对余额时。", "自动更新时间未到，但需要生成最新报表时。"):
        add_list_item(doc, text, bullet_id)

    add_page(doc, "12. 用户与邀请码管理（管理员）", "ADMINISTRATION", "“用户管理”仅对管理员显示。")
    add_heading(doc, "生成一次性邀请码", 2)
    for text in (
        "进入“用户管理”。",
        "在备注中填写邀请对象，便于后续识别。",
        "点击“生成邀请码”。",
        "立即复制并通过安全渠道发送；邀请码只显示这一次。",
    ):
        add_list_item(doc, text, number_id)
    add_heading(doc, "管理邀请码", 2)
    add_text(doc, "未使用的邀请码可以撤销；已经使用的邀请码不能再次使用，也不能作为新用户的注册凭证。")
    add_heading(doc, "管理已注册用户", 2)
    add_text(doc, "管理员可以查看注册用户，并按系统提供的操作启用或停用账号。停用前请确认影响范围。每位用户的负责人、手机、钱包、地址、币种和资产数据相互隔离。")
    add_callout(doc, "管理员安全建议", "不要在群聊或公开文档中长期保存邀请码；不再使用的账号应及时停用；管理员账号应使用独立的高强度密码。", "risk")

    add_page(doc, "13. 日常使用与安全规范", "OPERATIONS", "用统一命名和固定流程，能显著降低批量管理时的错误率。")
    add_heading(doc, "推荐命名", 2)
    add_table(doc, ["对象", "示例", "建议"], [
        ("负责人", "张三", "使用真实姓名或团队唯一代号"),
        ("手机", "ZS-P01", "负责人缩写 + 设备编号"),
        ("钱包", "OKX-01", "使用固定序号，避免“钱包1/主钱包”混用"),
        ("活动", "2026-08 活动 A", "包含时间或批次，便于归档"),
    ], [1800, 2500, 5060])
    add_heading(doc, "安全规范", 2)
    for text in (
        "系统只需要公开地址；绝不输入助记词、私钥、签名信息或交易所密码。",
        "自定义 RPC 必须来自可信来源，并优先使用 HTTPS。",
        "重要修改和批量删除前先核对负责人、手机和钱包归属。",
        "隐私模式适合屏幕共享，但导出 CSV 前仍要确认接收对象。",
        "不要把“女巫状态”当作系统自动结论；它是内部人工记录。",
    ):
        add_list_item(doc, text, bullet_id)
    add_heading(doc, "每日建议流程", 2)
    for text in ("查看总净值与昨日对比。", "检查待检查/疑似钱包数量。", "查看进行中活动的完成率。", "只对需要确认的负责人或钱包执行手动更新。", "必要时导出 CSV 留档。"):
        add_list_item(doc, text, number_id)

    add_page(doc, "14. 常见问题", "TROUBLESHOOTING", "遇到问题时，先检查归属、网络、报价和更新时间。")
    add_table(doc, ["问题", "原因与处理"], [
        ("导入地址后没有余额", "核对链和地址；执行手动更新；检查 RPC/网络是否可用。"),
        ("自定义币种识别成功但价值为 0", "通常是无有效 DEX 报价或流动性不足；编辑币种并填写手动单价。"),
        ("自定义币种只出现在一个钱包", "正常逻辑应为全账户扫描；先手动更新该链全部地址，并确认各钱包地址已正确导入。"),
        ("资产分布没有显示某币种", "该币种可能没有美元报价，或当前余额为零/被小额资产过滤。"),
        ("隐私模式仍看到金额", "先刷新页面后重新开启；金额应隐藏，但百分比和数量类统计可能保留。"),
        ("自动更新没有立即发生", "自动任务按 10–24 小时随机安排；需要即时数据请使用手动更新。"),
        ("页面仍是旧版本", "Windows 浏览器按 Ctrl+F5 强制刷新缓存。"),
        ("CSV 打开乱码", "优先用 Excel 的“从文本/CSV”导入，并选择 UTF-8 编码。"),
    ], [3100, 6260], font_size=9.2)
    add_callout(doc, "仍无法解决", "记录发生时间、页面名称、负责人/手机/钱包、链和公开地址，并截取错误提示。不要在截图或消息中发送任何私钥或助记词。", "info")

    add_page(doc, "15. 功能速查", "REFERENCE", "一页确认常用入口和操作结果。")
    add_table(doc, ["想做什么", "入口", "结果"], [
        ("查看全部资产", "资产总览", "总净值、时间对比、链/币种分布"),
        ("登记检查和活动", "钱包台账 / 钱包编辑", "状态统计、活动列和备注"),
        ("增加管理对象", "负责人 / 手机 / 钱包", "建立清晰的归属结构"),
        ("添加监控地址", "链上地址", "EVM 一次导入，多链复用"),
        ("补充非主流币", "资产明细 → 添加币种", "全账户识别、自动或手动报价"),
        ("接入新 EVM 链", "链配置 → 导入自定义链", "复用现有 EVM 地址"),
        ("立刻核对余额", "手动更新", "按全部/负责人/手机/钱包更新"),
        ("分享屏幕", "右上角隐私模式", "隐藏余额、数量与美元价值"),
        ("新增系统用户", "用户管理", "管理员生成一次性邀请码"),
    ], [2700, 2800, 3860], font_size=9.3)
    add_heading(doc, "术语", 2)
    add_table(doc, ["术语", "含义"], [
        ("WATCH ONLY", "只读监控，不具备签名和转账能力"),
        ("EVM 多链", "一套 0x 地址可在多条兼容链上查询"),
        ("自定义币种", "通过合约/Mint 地址补充识别的代币"),
        ("净值", "有有效价格的资产折算后的美元总价值"),
    ], [2300, 7060])
    add_text(doc, "— 文档结束 —", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
